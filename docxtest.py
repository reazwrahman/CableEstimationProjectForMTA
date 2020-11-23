#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Mon Jun 22 11:26:35 2020

@author: Reaz
"""

from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
from pandas.compat import BytesIO



class DOCX(object): 
    
    def __init__ (self,filename='Sample_Report.docx',heading='Report Shown Below'): 
        
        self.document=Document() 
        self.docname=filename
        self.heading=heading 
        self.document.add_heading(heading,0) 
        
        
        self.document.save(self.docname)
        
        
    def write_graph(self,graph_object,graph_name='Graph',image_width=5.25):  
    
        memfile = BytesIO() 
        graph_object.savefig(memfile)
 
        self.document.add_paragraph(graph_name,style='List Bullet')
        self.document.add_picture(memfile,width=Inches(image_width))  
        self.document.save(self.docname) 
        memfile.close()
        
        
    








#document.add_paragraph(
#    'Graph1', style='List Bullet')
#document.add_picture(memfile, width=Inches(5.25))
#
#document.add_paragraph(
#    'Graph2', style='List Number')
#document.add_picture(memfile, width=Inches(5.25))




#def write_image(): 
#    memfile = BytesIO() 
#    t = np.arange(0.0ß, 2.0, 0.01)
#    s = 1 + np.sin(2*np.pi*t)
#    plt.plot(t, s)
#    plt.savefig(memfile) 
#    document.add_paragraph(
#    'Graph1', style='List Bullet')
#    document.add_picture(memfile, width=Inches(5.25))  
#    document.save('report.docx')
#    memfile.close()
#
#write_image() 
#write_image() 
#write_image()
#

#document.save('report.docx')
#memfile.close() 
        
